#!/usr/bin/env python3
"""Test script for v0.2.3 improvements."""

import sys
from pathlib import Path
from report_killer.agent import ReportAgent
from report_killer.config import Config
from report_killer.docx_handler import DocxHandler, InsertionPoint
from rich.console import Console

console = Console()

def test_markdown_table_parsing():
    """Test Markdown table parsing."""
    console.print("\n[bold cyan]Test 1: Markdown Table Parsing[/bold cyan]")
    
    config = Config()
    agent = ReportAgent(config, test_mode=True)
    
    # Test table
    table_text = """| Name | Age | City |
|------|-----|------|
| Alice | 25 | NYC |
| Bob | 30 | LA |
| Charlie | 35 | SF |"""
    
    result = agent._parse_markdown_table(table_text)
    if result:
        headers, rows = result
        console.print(f"  Headers: {headers}")
        console.print(f"  Rows: {len(rows)} rows")
        console.print(f"  [green]✓ Table parsing works[/green]")
        return True
    else:
        console.print(f"  [red]✗ Table parsing failed[/red]")
        return False

def test_config_max_retries():
    """Test that max_code_retries is in config."""
    console.print("\n[bold cyan]Test 2: Config Max Retries[/bold cyan]")
    
    config = Config()
    if hasattr(config, 'max_code_retries'):
        console.print(f"  max_code_retries: {config.max_code_retries}")
        console.print(f"  [green]✓ Config has max_code_retries[/green]")
        return True
    else:
        console.print(f"  [red]✗ Config missing max_code_retries[/red]")
        return False

def test_code_block_no_label():
    """Test that code blocks don't have language labels."""
    console.print("\n[bold cyan]Test 3: Code Block Without Language Label[/bold cyan]")
    
    # Create a test document
    test_doc_path = Path("workspace/test_code_label.docx")
    test_doc_path.parent.mkdir(exist_ok=True)
    
    from docx import Document
    doc = Document()
    doc.add_paragraph("Test document")
    doc.add_paragraph("Question: What is Python?")
    doc.save(test_doc_path)
    
    # Load and insert code
    handler = DocxHandler(str(test_doc_path))
    handler.load()
    
    code = "print('Hello, World!')\nprint('Line 2')"
    num_inserted = handler.insert_code_block(1, code, "python")
    
    # Check that we only inserted code lines, not a label
    # Code has 2 lines, so should insert 2 paragraphs
    if num_inserted == 2:
        console.print(f"  Inserted {num_inserted} paragraphs (code lines only)")
        console.print(f"  [green]✓ No language label inserted[/green]")
        
        # Verify paragraphs
        handler.save(test_doc_path)
        handler2 = DocxHandler(str(test_doc_path))
        handler2.load()
        
        # Should have: original para + question + 2 code lines = 4 total
        total = len([p for p in handler2.doc.paragraphs if p.text.strip()])
        console.print(f"  Total paragraphs with text: {total}")
        
        test_doc_path.unlink()
        return True
    else:
        console.print(f"  [red]✗ Expected 2 paragraphs, got {num_inserted}[/red]")
        return False

def test_table_insertion():
    """Test table insertion from parsed markdown."""
    console.print("\n[bold cyan]Test 4: Table Insertion in Document[/bold cyan]")
    
    test_doc_path = Path("workspace/test_table.docx")
    test_doc_path.parent.mkdir(exist_ok=True)
    
    from docx import Document
    doc = Document()
    doc.add_paragraph("Before table")
    doc.add_paragraph("After table")
    doc.save(test_doc_path)
    
    # Insert table
    handler = DocxHandler(str(test_doc_path))
    handler.load()
    
    headers = ["Name", "Score"]
    data = [["Alice", "95"], ["Bob", "87"]]
    
    num_inserted = handler.insert_table(0, data, headers)
    handler.save(test_doc_path)
    
    # Reload and verify
    handler2 = DocxHandler(str(test_doc_path))
    handler2.load()
    
    num_tables = len(handler2.doc.tables)
    if num_tables == 1:
        table = handler2.doc.tables[0]
        console.print(f"  Table rows: {len(table.rows)}, cols: {len(table.columns)}")
        console.print(f"  [green]✓ Table inserted successfully[/green]")
        
        test_doc_path.unlink()
        return True
    else:
        console.print(f"  [red]✗ Expected 1 table, found {num_tables}[/red]")
        return False

def test_response_parsing_with_table():
    """Test that response with markdown table gets parsed correctly."""
    console.print("\n[bold cyan]Test 5: Response Parsing with Markdown Table[/bold cyan]")
    
    config = Config()
    agent = ReportAgent(config, test_mode=True)
    
    # Create a fake insertion point
    point = InsertionPoint(10, "Test question", "", "")
    
    # Response with markdown table
    response = """Here is the analysis:

| Algorithm | Time | Space |
|-----------|------|-------|
| BFS | O(n) | O(n) |
| DFS | O(n) | O(h) |

The table shows the comparison."""
    
    content = agent._parse_response_to_structured_content(response, point)
    
    # Check that we got text, table, text
    types = [item['type'] for item in content]
    console.print(f"  Content types: {types}")
    
    if 'table' in types:
        table_items = [item for item in content if item['type'] == 'table']
        if table_items:
            table = table_items[0]
            console.print(f"  Table headers: {table.get('headers')}")
            console.print(f"  Table rows: {len(table.get('data', []))}")
            console.print(f"  [green]✓ Markdown table parsed from response[/green]")
            return True
    
    console.print(f"  [red]✗ No table found in parsed content[/red]")
    return False

def main():
    console.print("[bold]Testing v0.2.3 Improvements[/bold]")
    console.print("=" * 60)
    
    tests = [
        test_markdown_table_parsing,
        test_config_max_retries,
        test_code_block_no_label,
        test_table_insertion,
        test_response_parsing_with_table,
    ]
    
    results = []
    for test in tests:
        try:
            result = test()
            results.append(result)
        except Exception as e:
            console.print(f"  [red]✗ Test failed with exception: {e}[/red]")
            results.append(False)
    
    console.print("\n" + "=" * 60)
    passed = sum(results)
    total = len(results)
    console.print(f"[bold]Results: {passed}/{total} tests passed[/bold]")
    
    if passed == total:
        console.print("[green]✓ All tests passed![/green]")
        return 0
    else:
        console.print(f"[red]✗ {total - passed} test(s) failed[/red]")
        return 1

if __name__ == "__main__":
    sys.exit(main())
