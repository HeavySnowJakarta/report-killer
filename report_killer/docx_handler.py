"""Word document handling utilities using python-docx."""

from typing import List, Dict, Optional, Tuple
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
from pathlib import Path


class InsertionPoint:
    """Represents a point in the document where content should be inserted."""
    
    def __init__(self, para_index: int, description: str, context_before: str = "", context_after: str = ""):
        self.para_index = para_index
        self.description = description
        self.context_before = context_before
        self.context_after = context_after
        self.filled = False
        self.content = []
        self.actual_insert_index = para_index  # Track where we actually inserted
    
    def __repr__(self):
        status = "✓" if self.filled else "○"
        return f"{status} Para {self.para_index}: {self.description[:50]}..."


class DocxHandler:
    """Handle reading and writing Word documents using python-docx."""
    
    def __init__(self, filepath: str):
        """Initialize with a document filepath."""
        self.filepath = Path(filepath)
        self.doc = None
        self.insertion_points: List[InsertionPoint] = []
        
    def load(self):
        """Load the document."""
        self.doc = Document(self.filepath)
        
    def get_text_content(self) -> str:
        """Extract all text content from the document."""
        if not self.doc:
            self.load()
        
        content = []
        for para in self.doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # Also extract text from tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        content.append(cell.text)
        
        return "\n".join(content)
    
    def get_paragraphs_with_indices(self) -> List[Tuple[int, str]]:
        """Get all paragraphs with their indices."""
        if not self.doc:
            self.load()
        
        result = []
        for idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if text:
                result.append((idx, text))
        
        return result
    
    def get_context_around_index(self, para_index: int, before: int = 5, after: int = 5) -> Tuple[str, str]:
        """Get context before and after a paragraph index."""
        if not self.doc:
            self.load()
        
        before_text = []
        after_text = []
        
        for i in range(max(0, para_index - before), para_index):
            if i < len(self.doc.paragraphs):
                text = self.doc.paragraphs[i].text.strip()
                if text:
                    before_text.append(text)
        
        for i in range(para_index + 1, min(len(self.doc.paragraphs), para_index + after + 1)):
            text = self.doc.paragraphs[i].text.strip()
            if text:
                after_text.append(text)
        
        return "\n".join(before_text), "\n".join(after_text)
    
    def set_insertion_points(self, points: List[InsertionPoint]):
        """Set the insertion points from LLM analysis."""
        self.insertion_points = points
    
    def insert_paragraph_after(self, para_index: int, text: str, style: Optional[str] = None):
        """Insert a new paragraph after the specified paragraph."""
        if not self.doc:
            self.load()
        
        if para_index >= len(self.doc.paragraphs):
            # Append at end
            new_para = self.doc.add_paragraph(text)
        else:
            # Get the paragraph to insert after
            target_para = self.doc.paragraphs[para_index]
            
            # Insert after by getting the parent element and inserting
            new_para = target_para.insert_paragraph_before(text)
            
            # Move it to after the target
            target_para._element.addnext(new_para._element)
        
        if style:
            new_para.style = style
        
        return new_para
    
    def insert_code_block(self, para_index: int, code: str, language: str = "") -> int:
        """Insert a code block with proper formatting. Returns number of paragraphs inserted."""
        if not self.doc:
            self.load()
        
        inserted = 0
        
        # Split code into lines to preserve line breaks
        code_lines = code.split('\n')
        
        # Insert code with line breaks preserved (no language label)
        for line in code_lines:
            code_para = self.insert_paragraph_after(para_index + inserted, line if line.strip() else " ")
            
            # Format as code
            for run in code_para.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                # Use monospace font for Chinese characters too
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            
            inserted += 1
        
        return inserted
    
    def insert_table(self, para_index: int, data: List[List[str]], headers: Optional[List[str]] = None) -> int:
        """Insert a table after the specified paragraph. Returns 1."""
        if not self.doc:
            self.load()
        
        rows = len(data) + (1 if headers else 0)
        cols = len(headers) if headers else len(data[0]) if data else 0
        
        if rows == 0 or cols == 0:
            return 0
        
        # Create table
        if para_index >= len(self.doc.paragraphs):
            table = self.doc.add_table(rows=rows, cols=cols)
        else:
            # Insert table after paragraph
            target_para = self.doc.paragraphs[para_index]
            table = self.doc.add_table(rows=rows, cols=cols)
            # Move table to correct position
            target_para._element.addnext(table._element)
        
        table.style = 'Light Grid Accent 1'
        
        # Fill headers
        if headers:
            for col_idx, header in enumerate(headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = header
                # Make header bold
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Fill data
        start_row = 1 if headers else 0
        for row_idx, row_data in enumerate(data):
            for col_idx, cell_data in enumerate(row_data):
                table.rows[start_row + row_idx].cells[col_idx].text = str(cell_data)
        
        return 1  # Table counts as 1 element
    
    def insert_image(self, para_index: int, image_path: str, width_inches: float = 5.0) -> int:
        """Insert an image after the specified paragraph. Returns 1."""
        if not self.doc:
            self.load()
        
        if para_index >= len(self.doc.paragraphs):
            para = self.doc.add_paragraph()
        else:
            target_para = self.doc.paragraphs[para_index]
            para = target_para.insert_paragraph_before()
            target_para._element.addnext(para._element)
        
        run = para.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))
        
        return 1
    
    def insert_content_at_point(self, point: InsertionPoint, content: List[dict]):
        """
        Insert content at the specified insertion point.
        
        Content is a list of dicts with:
        - type: 'text', 'code', 'table', 'image'
        - data: the actual content
        - language: (for code) programming language
        - headers: (for table) optional headers
        """
        if not self.doc:
            self.load()
        
        current_index = point.para_index
        inserted_total = 0
        
        for item in content:
            item_type = item.get('type', 'text')
            
            if item_type == 'text':
                text = item.get('data', '')
                self.insert_paragraph_after(current_index, text)
                inserted_total += 1
                current_index += 1
                
            elif item_type == 'code':
                code = item.get('data', '')
                language = item.get('language', '')
                inserted = self.insert_code_block(current_index, code, language)
                inserted_total += inserted
                current_index += inserted
                
            elif item_type == 'table':
                data = item.get('data', [])
                headers = item.get('headers')
                inserted = self.insert_table(current_index, data, headers)
                inserted_total += inserted
                current_index += inserted
                
            elif item_type == 'image':
                image_path = item.get('data', '')
                if Path(image_path).exists():
                    width = item.get('width', 5.0)
                    inserted = self.insert_image(current_index, image_path, width)
                    inserted_total += inserted
                    current_index += inserted
        
        point.filled = True
        point.content = content
        point.actual_insert_index = current_index
        
        # Update subsequent insertion points' indices
        for other_point in self.insertion_points:
            if other_point.para_index > point.para_index and not other_point.filled:
                other_point.para_index += inserted_total
    
    def save(self, output_path: Optional[str] = None):
        """Save the document."""
        if not self.doc:
            raise ValueError("No document loaded")
        
        save_path = output_path or str(self.filepath)
        self.doc.save(save_path)
    
    def get_completion_status(self) -> Dict:
        """Get status of all insertion points."""
        total = len(self.insertion_points)
        filled = sum(1 for p in self.insertion_points if p.filled)
        
        return {
            "total": total,
            "filled": filled,
            "remaining": total - filled,
            "completion_rate": filled / total if total > 0 else 0,
            "points": self.insertion_points,
        }
